from docx import Document
import io

def extract_brd():
    try:
        doc = Document(r'd:\code\GI\reports\FIS_Graph_Intelligence_-_BRD_V0.2_DRAFT.docx')
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract tables if any
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                full_text.append(" | ".join(cells))

        with open('brd_v02_extracted.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        print("Successfully extracted BRD to brd_extracted.txt")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_brd()
